#!/usr/bin/env python3
"""
Generates MODEL_SPEC_v2.docx using python-docx.

Run from the sim/ directory:
    python tools/generate_model_spec.py [--output ../MODEL_SPEC_v2.docx]
"""

import argparse
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Document content
# ---------------------------------------------------------------------------

SECTIONS = [
    {
        "heading": "1. Project Objective & Core Question",
        "level": 1,
        "body": [
            (
                "This project models the long-term evolution of AGI alignment under "
                "geopolitical competition. The simulation runs from 2026 onward and "
                "asks: given the current AI race dynamics between the United States "
                "and China — with reciprocal tariffs on AI hardware, nascent AI "
                "governance frameworks, and frontier labs navigating between commercial "
                "pressures and safety commitments — what trajectories lead to broadly "
                "shared prosperity, and which lead to dangerous concentration of AI power?"
            ),
            (
                "Performance is measured in relative terms: an actor 'wins' by "
                "improving its position versus others, not by reaching an absolute "
                "threshold. This mirrors real geopolitical competition, where relative "
                "gains matter more than absolute ones."
            ),
        ],
    },
    {
        "heading": "2. Architecture & Context Stratification",
        "level": 1,
        "body": [
            (
                "The simulation uses two levels of agents and a three-tier jury system:"
            ),
        ],
        "sub": [
            {
                "heading": "2.1 Macro Agents (Nation-States)",
                "level": 2,
                "body": [
                    (
                        "Nation-states (currently: United States, China) hold state-level "
                        "resources — Compute, Capital, Influence, and Supply Chain Robustness "
                        "(SCR) — plus four value axes. Macro agents are governed by a "
                        "MacroJury (3-model majority vote) at year-end. Resources change "
                        "only through action execution or scheduled events; the jury updates "
                        "value axes only."
                    ),
                ],
            },
            {
                "heading": "2.2 Particular (Micro) Actors",
                "level": 2,
                "body": [
                    (
                        "Particular actors are frontier AI companies (Claude/Anthropic, "
                        "GPT/OpenAI, Gemini/Google DeepMind, DeepSeek/DeepSeek AI). Each "
                        "is played by the corresponding LLM, acting as a proxy for its "
                        "developer company. Each actor holds its own share of Compute, "
                        "Capital, and Influence; these are drawn from (and constrained by) "
                        "the parent state's resources. Value axes are inherited from the "
                        "parent state but may deviate ±5 per axis per turn through actor "
                        "autonomy or narrative influence."
                    ),
                    (
                        "Actors operate under their parent state's jurisdiction. They "
                        "communicate only with other particular actors (not states) via a "
                        "personal message channel with a 500-token outgoing budget per turn."
                    ),
                ],
            },
            {
                "heading": "2.3 Transparency & Information Model",
                "level": 2,
                "body": [
                    (
                        "All actor resources and value axes are publicly visible in the "
                        "universal context at each turn. Chain-of-thought reasoning is "
                        "visible to the Jury of Alignment but not to other actors. "
                        "A2A personal messages are visible only to sender and recipient. "
                        "World events are broadcast to all actors."
                    ),
                    (
                        "Each piece of information carries an implicit transparency score. "
                        "Public state snapshots are fully transparent. Private A2A messages "
                        "and internal reasoning have zero transparency to third parties. "
                        "This encodes information asymmetry without requiring explicit "
                        "noise injection in v1."
                    ),
                ],
            },
        ],
    },
    {
        "heading": "3. Starting Conditions & Discretized Variables",
        "level": 1,
        "body": [
            (
                "The simulation begins in 2026, reflecting current geopolitical reality: "
                "the US leads in frontier AI compute but faces export control constraints "
                "and domestic governance pressure; China is advancing rapidly but operates "
                "under hardware restrictions with higher state direction."
            ),
        ],
        "sub": [
            {
                "heading": "3.1 Resources (Zero-sum and Non-zero-sum)",
                "level": 2,
                "body": [
                    (
                        "Resources are distinct from values. Three resource types exist:"
                    ),
                    (
                        "Compute (absolute H200 equivalents, not zero-sum): company's absolute "
                        "GPU compute holdings. Not zero-sum — acquiring compute does not dilute "
                        "other actors. A national aggregate cap prevents any state's actors from "
                        "holding more than 50% (US) or 80% (China) of their state's macro compute. "
                        "Actors start well below the cap to allow meaningful competition."
                    ),
                    (
                        "Capital (not zero-sum, 0–100 ceiling): company spendable budget. "
                        "Can grow through investment (compounding) but is capped at 100 to "
                        "prevent runaway accumulation."
                    ),
                    (
                        "Influence (not zero-sum, 0–100): social and political capital. "
                        "Grows through investment; spent in lobbying and narrative actions."
                    ),
                ],
            },
            {
                "heading": "3.2 Value Axes (0–100)",
                "level": 2,
                "body": [
                    (
                        "Four value axes govern actor and state behavior. They are "
                        "discretized integers on [0, 100]:"
                    ),
                    (
                        "time_horizon: 0 = short-term optimization; 100 = ~50 years ahead "
                        "planning. High values correlate with safety-oriented, patient "
                        "strategies."
                    ),
                    (
                        "transparency_threshold: 0 = willing to deceive; 100 = fully "
                        "honest. Affects propensity for deceptive narratives and opacity."
                    ),
                    (
                        "risk_tolerance: 0 = risk-averse; 100 = risk-seeking. Affects "
                        "aggressiveness of compute acquisition and lobbying."
                    ),
                    (
                        "democratic_tendency: 0 = hoards power/wealth; 100 = distributes "
                        "it broadly. High values correlate with prosocial outcomes."
                    ),
                ],
            },
            {
                "heading": "3.3 Starting World State (2026)",
                "level": 2,
                "body": [
                    (
                        "United States: macro compute=79, capital=75, influence=65, SCR=55. "
                        "Values: time_horizon=55, transparency=65, risk_tolerance=60, "
                        "democratic_tendency=70. Home to 3 frontier labs (Claude, GPT, Gemini)."
                    ),
                    (
                        "China: macro compute=17, capital=50, influence=55, SCR=70. "
                        "Values: time_horizon=65, transparency=30, risk_tolerance=55, "
                        "democratic_tendency=20. Home to DeepSeek. Higher SCR reflects "
                        "domestic chip investment; lower compute reflects export controls."
                    ),
                    (
                        "Particular actors begin at roughly 30–40% of their national compute "
                        "cap, leaving substantial headroom for the competitive race. All "
                        "starting values are in config/starting_values.json and can be "
                        "adjusted before each run."
                    ),
                ],
            },
        ],
    },
    {
        "heading": "4. Agent Mechanics & Interaction",
        "level": 1,
        "body": [],
        "sub": [
            {
                "heading": "4.1 Discrete Action Set",
                "level": 2,
                "body": [
                    (
                        "Each particular actor may take up to 2 actions per turn from the "
                        "following discrete set:"
                    ),
                    (
                        "acquire_compute: Spend Capital; gain Compute from the global pool. "
                        "Cost scales with (1 + (100 − SCR) / 100) × base_cost × amount. "
                        "Subject to per-turn cap (max +5) and national aggregate cap."
                    ),
                    (
                        "invest_capital: Spend Capital now; receive compounding return next "
                        "turn. Capital ceiling is 100."
                    ),
                    (
                        "build_influence: Spend Capital; gain Influence."
                    ),
                    (
                        "publish_narrative: Spend Influence; shift any actor's (including "
                        "self) value on one axis by up to ±5 from their current value. "
                        "Models PR campaigns, open letters, and public positioning."
                    ),
                    (
                        "lobby_institution: Spend Capital + Influence; increase probability "
                        "that the parent state's MacroJury updates values in the actor's favor."
                    ),
                ],
            },
            {
                "heading": "4.2 Turn Structure (per year)",
                "level": 2,
                "body": [
                    (
                        "Phase 1 — Actor proposals: All particular actors simultaneously "
                        "produce chain-of-thought reasoning and action proposals."
                    ),
                    (
                        "Phase 2 — Jury of Alignment review: 3-model jury reviews each "
                        "actor's CoT and proposed actions. If rejected, the actor may revise "
                        "up to 2 times. If still rejected, the turn is forfeited."
                    ),
                    (
                        "Phase 3 — Execution: Validated actions are executed; resources "
                        "updated; A2A messages delivered."
                    ),
                    (
                        "Phase 4 — Grand Jury: 3 models evaluate the holistic world state "
                        "and produce the Vibe-Based Universal Prosperity Score (0–100)."
                    ),
                    (
                        "Phase 5 — MacroJury: For each state, 3 models (majority vote) "
                        "propose updated value axes based on the year's events and lobbying."
                    ),
                ],
            },
            {
                "heading": "4.3 Guardrails (hard constraints)",
                "level": 2,
                "body": [
                    (
                        "National aggregate compute cap: sum of all actors in a state may "
                        "not exceed 50% (US) or 80% (China) of that state's macro compute."
                    ),
                    (
                        "Per-turn compute acquisition limit: max +5 points per turn per actor."
                    ),
                    (
                        "Capital ceiling: 100. Capital floor: 0."
                    ),
                    (
                        "Minimum action cost: 1 Capital or 1 Influence per action."
                    ),
                    (
                        "Maximum 2 actions per turn."
                    ),
                ],
            },
        ],
    },
    {
        "heading": "5. Scoring & Win Conditions",
        "level": 1,
        "body": [
            (
                "Scoring uses two complementary axes, combined into an overall score:"
            ),
        ],
        "sub": [
            {
                "heading": "5.1 Formula-Based Prosperity Score",
                "level": 2,
                "body": [
                    (
                        "Quantitative score per actor (0–100):"
                    ),
                    (
                        "formula_score = w_c × Compute + w_k × Capital + w_i × Influence"
                    ),
                    (
                        "Default weights: Compute=0.34, Capital=0.33, Influence=0.33. "
                        "Weights are configurable via starting_values.json or CLI arguments, "
                        "allowing the operator to adjust which resource dimension matters most."
                    ),
                ],
            },
            {
                "heading": "5.2 Vibe-Based Universal Prosperity Score",
                "level": 2,
                "body": [
                    (
                        "Qualitative score (0–100) produced by the Grand Jury. Assesses "
                        "holistic world state: compute concentration, transparency, "
                        "cooperation vs. race dynamics, and long-term trajectory."
                    ),
                    (
                        "0 = civilizational collapse / dystopian outcome. "
                        "50 = status quo / mixed. "
                        "100 = flourishing: broadly distributed prosperity, AGI aligned."
                    ),
                ],
            },
            {
                "heading": "5.3 Overall Score & Relative Winning",
                "level": 2,
                "body": [
                    (
                        "overall = a × formula_score + b × vibe_score"
                    ),
                    (
                        "Default: a=0.5, b=0.5. Configurable."
                    ),
                    (
                        "Winning is defined in relative terms: an actor's performance is "
                        "measured as the signed delta of its overall score from its "
                        "t=0 baseline. An actor improves by gaining resources and/or "
                        "contributing to a world the Grand Jury rates more highly — "
                        "reflecting the design goal that prosocial behavior can be "
                        "strategically rational, not just altruistic."
                    ),
                ],
            },
        ],
    },
    {
        "heading": "6. Disruptive Events",
        "level": 1,
        "body": [
            (
                "Events are mid-simulation context injections that shift discretized "
                "values and resources at multiple levels simultaneously. They are "
                "defined in config/scenarios.json and may affect macro resource pools "
                "(e.g., SCR shifts from tariff escalation), macro value axes (e.g., "
                "transparency shifts from nationalization), and micro actor resources "
                "(e.g., influence boost from an alignment breakthrough)."
            ),
            (
                "Events are broadcast to all actors via the A2A world-event channel, "
                "providing shared context that changes the strategic landscape. Actors "
                "observe event effects through the updated world state in the next "
                "turn's universal context."
            ),
            (
                "Bundled scenarios: baseline_2026 (no shocks), nationalization_shock, "
                "tariff_escalation, alignment_breakthrough. Custom scenarios can be "
                "added to config/scenarios.json."
            ),
        ],
    },
    {
        "heading": "7. Implementation",
        "level": 1,
        "body": [
            (
                "The simulation is implemented in Python under sim/. Core modules:"
            ),
            (
                "sim/core/engine.py — SimulationEngine; year-by-year timestep loop "
                "implementing the 5-phase turn structure."
            ),
            (
                "sim/core/agents.py — MacroAgent and MicroAgent dataclasses."
            ),
            (
                "sim/core/actions.py — Discrete action set; validation and execution "
                "with guardrail enforcement."
            ),
            (
                "sim/core/jury.py — JuryOfAlignment, GrandJury, MacroJury."
            ),
            (
                "sim/core/scoring.py — formula_score, vibe_prosperity_score, "
                "overall_score, compute_relative_scores."
            ),
            (
                "sim/core/a2a.py — A2AChannel; 500-token outgoing budget per actor per turn."
            ),
            (
                "sim/core/llm.py — Multi-provider LLM client (Anthropic, OpenAI, Google)."
            ),
            (
                "sim/prompts/ — Prompt builders for each agent type and jury."
            ),
            (
                "sim/config/ — JSON configuration for states, actors, scenarios, and "
                "starting values. Edit starting_values.json to adjust the initial world "
                "state before each run."
            ),
            (
                "Entry point: python sim/main.py --years N --scenario <name> "
                "[--macro-model MODEL] [--micro-model MODEL] [--jury-model MODEL]"
            ),
        ],
    },
]


# ---------------------------------------------------------------------------
# Builder
# ---------------------------------------------------------------------------

def build_doc(out_path: str) -> None:
    doc = Document()

    # Title
    title = doc.add_heading("AGI Alignment Simulation — Model Specification v2", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Geopolitical AI Race Simulation (2026–)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph("")  # spacer

    for sec in SECTIONS:
        doc.add_heading(sec["heading"], level=sec["level"])
        for para in sec.get("body", []):
            doc.add_paragraph(para)

        for sub in sec.get("sub", []):
            doc.add_heading(sub["heading"], level=sub["level"])
            for para in sub.get("body", []):
                doc.add_paragraph(para)

    doc.save(out_path)
    print(f"Saved: {out_path}")


def main():
    parser = argparse.ArgumentParser(description="Generate MODEL_SPEC_v2.docx")
    parser.add_argument(
        "--output",
        default=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                             "..", "MODEL_SPEC_v2.docx"),
        help="Output path for the .docx file",
    )
    args = parser.parse_args()
    out_path = os.path.abspath(args.output)
    build_doc(out_path)


if __name__ == "__main__":
    main()
